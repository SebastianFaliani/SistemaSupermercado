from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / 'documentacion' / 'Guia_de_Prueba_Funcional_La_91_v0.2.docx'
LOGO = ROOT / 'frontend' / 'public' / 'marca' / 'logo-horizontal-claro.png'
AZUL, VERDE, CELESTE, CLARO, GRIS = '003B46', '07575B', '66A5AD', 'C4DFE6', '52656A'

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.49)

def aplicar_fuente(run, size=10.5, color=AZUL, bold=False, italic=False):
    run.font.name = 'Aptos'
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn('w:ascii'), 'Aptos'); rfonts.set(qn('w:hAnsi'), 'Aptos')
    run.font.size = Pt(size); run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold; run.italic = italic
    return run

normal = doc.styles['Normal']; normal.font.name = 'Aptos'; normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string(AZUL)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.15
for name, size, before, after, color in [('Title', 30, 0, 8, AZUL), ('Subtitle', 14, 0, 8, VERDE), ('Heading 1', 18, 18, 9, AZUL), ('Heading 2', 14, 14, 7, VERDE), ('Heading 3', 11.5, 10, 5, AZUL)]:
    estilo = doc.styles[name]; estilo.font.name = 'Aptos'; estilo.font.size = Pt(size); estilo.font.bold = name != 'Subtitle'; estilo.font.color.rgb = RGBColor.from_string(color)
    estilo.paragraph_format.space_before = Pt(before); estilo.paragraph_format.space_after = Pt(after); estilo.paragraph_format.keep_with_next = True
for name in ['List Bullet', 'List Number']:
    estilo = doc.styles[name]; estilo.font.name = 'Aptos'; estilo.font.size = Pt(10.5)
    estilo.paragraph_format.left_indent = Inches(.38); estilo.paragraph_format.first_line_indent = Inches(-.19); estilo.paragraph_format.space_after = Pt(4); estilo.paragraph_format.line_spacing = 1.15

def sombrear(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = tc_pr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)

def margen_celda(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr(); tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None: tc_mar = OxmlElement('w:tcMar'); tc_pr.append(tc_mar)
    for nombre, valor in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        nodo = tc_mar.find(qn(f'w:{nombre}'))
        if nodo is None: nodo = OxmlElement(f'w:{nombre}'); tc_mar.append(nodo)
        nodo.set(qn('w:w'), str(valor)); nodo.set(qn('w:type'), 'dxa')

def tabla(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False; t.style = 'Table Grid'
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]; cell.text = ''; sombrear(cell, AZUL); margen_celda(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        aplicar_fuente(cell.paragraphs[0].add_run(header), 9.3, 'FFFFFF', True)
    tr_pr = t.rows[0]._tr.get_or_add_trPr(); repetir = OxmlElement('w:tblHeader'); repetir.set(qn('w:val'), 'true'); tr_pr.append(repetir)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ''; margen_celda(cells[i]); cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            aplicar_fuente(cells[i].paragraphs[0].add_run(str(value)), 9.1, AZUL)
    for row in t.rows:
        for i, width in enumerate(widths): row.cells[i].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def nota(titulo, texto, color=CLARO):
    t = doc.add_table(rows=1, cols=1); t.autofit = False; t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.columns[0].width = Inches(6.5)
    cell = t.cell(0, 0); sombrear(cell, color); margen_celda(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]; aplicar_fuente(p.add_run(titulo + ' '), 10, AZUL, True); aplicar_fuente(p.add_run(texto), 10, AZUL)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def heading(texto, nivel=1): doc.add_heading(texto, level=nivel)
def paso(texto): doc.add_paragraph(texto, style='List Number')
def item(texto): doc.add_paragraph(texto, style='List Bullet')
def check(texto): doc.add_paragraph('[ ] ' + texto)
def pagina(): doc.add_page_break()

header = sec.header.paragraphs[0]; header.alignment = WD_ALIGN_PARAGRAPH.LEFT; aplicar_fuente(header.add_run('LA 91 SUPERMERCADO  |  GUÍA DE PRUEBA FUNCIONAL'), 8.5, VERDE, True)
footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; aplicar_fuente(footer.add_run('Versión incremental 0.2 · Agosto 2026 · Capacitación y aceptación'), 8, GRIS)

doc.add_paragraph().paragraph_format.space_after = Pt(60)
if LOGO.exists():
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imagen = p.add_run().add_picture(str(LOGO), width=Inches(4.9)); imagen._inline.docPr.set('descr', 'Logotipo de La 91 Supermercado')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; aplicar_fuente(p.add_run('GUÍA DE PRUEBA FUNCIONAL'), 12, CELESTE, True)
p = doc.add_paragraph(); p.style = 'Title'; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Sistema de Gestión\nLa 91 Supermercado')
p = doc.add_paragraph(); p.style = 'Subtitle'; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Recorrido paso a paso con datos de ejemplo y resultados esperados')
doc.add_paragraph().paragraph_format.space_after = Pt(70)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; aplicar_fuente(p.add_run('VERSIÓN INCREMENTAL 0.2'), 11, VERDE, True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; aplicar_fuente(p.add_run('Actualizada al 10 de agosto de 2026'), 10, GRIS)
nota('Documento vivo.', 'Esta guía registra las pruebas realizadas y se ampliará después de cada nuevo circuito validado. No contiene contraseñas reales.', 'EAF4F6')

pagina(); heading('1. Objetivo y forma de uso')
doc.add_paragraph('Esta guía permite que el cliente recorra el sistema desde una base operativa controlada, comprenda qué impacto produce cada acción y confirme que los saldos coincidan. Los importes son ejemplos de capacitación y deben reemplazarse por datos reales al poner el sistema en producción.')
nota('Regla de prueba.', 'Ejecute los pasos en el orden indicado. Antes de continuar, marque cada control como correcto. Si un importe no coincide, deténgase y registre la incidencia; no repita la operación ni cree un movimiento manual para compensarla.')
heading('1.1 Alcance de esta edición', 2)
for texto in ['Acceso, usuarios y cajas.', 'Tesorería inicial.', 'Clientes, proveedores y empleados de prueba.', 'Inventario y ajustes iniciales.', 'Orden de compra y recepción parcial.', 'Actualización automática de costo y precio de venta.', 'Factura y cuenta corriente de proveedor.', 'Pago parcial vinculado con Tesorería.', 'Gastos recurrentes, edición, anulación y pago parcial.']: item(texto)
heading('1.2 Convenciones', 2)
tabla(['Elemento', 'Significado'], [('Acción', 'Paso que debe ejecutar el operador.'), ('Resultado esperado', 'Dato que debe mostrar el sistema después de guardar.'), ('Control', 'Verificación que debe marcarse antes de continuar.'), ('Detenerse', 'No continuar hasta resolver una diferencia.')], [1.5, 5.0])

heading('2. Datos utilizados en la prueba')
tabla(['Tipo', 'Dato de ejemplo'], [('Usuario administrador', 'administrador · Sebastian Faliani'), ('Correo', 'nedase@gmail.com'), ('Cliente', 'Cliente Prueba Cuenta Corriente · documento 99999992'), ('Proveedor', 'Distribuidora Prueba'), ('Empleado', 'PRUEBA-001 · documento 99999993 · modalidad quincenal'), ('Cuenta bancaria', 'Banco de prueba'), ('Cuenta de efectivo', 'Efectivo general'), ('Orden de compra', '#1 · 6 productos · total $303.956,98'), ('Factura proveedor', 'Factura A-0001-00000001 · $303.956,98')], [1.75, 4.75])
nota('Seguridad.', 'La contraseña de prueba no se escribe en este documento. Cada usuario debe utilizar una clave personal de al menos 12 caracteres.')

pagina(); heading('3. Preparación y acceso')
heading('3.1 Iniciar el sistema', 2)
paso('Abra una terminal en la carpeta del sistema.')
paso('Ejecute npm run dev.')
paso('Abra http://localhost:5173 en el navegador.')
paso('Inicie sesión con un usuario autorizado.')
for texto in ['El backend informa API disponible en http://localhost:3000.', 'La pantalla no muestra errores de conexión.', 'Las fechas visibles utilizan dd-mm-aaaa.', 'Las horas utilizan formato de 24 horas, de 00:00 a 23:59.']: check(texto)
heading('3.2 Usuarios y contraseñas', 2)
paso('Abra Usuarios y seleccione Nuevo usuario o Editar.')
paso('Escriba la contraseña y utilice el icono del ojo para mostrar u ocultar el contenido.')
paso('Compruebe el mismo comportamiento en todos los formularios que soliciten contraseña.')
heading('3.3 Cajas', 2)
paso('Abra Punto de venta y acceda a Administrar cajas.')
paso('Cree o edite una caja física con código y nombre identificables.')
paso('Abra una sesión seleccionando una caja disponible e ingresando el fondo inicial contado.')
nota('Control de concurrencia.', 'Una caja física no debe quedar abierta para dos sesiones. Cada operador utiliza su propio usuario y selecciona una caja disponible.')

pagina(); heading('4. Tesorería inicial')
heading('4.1 Crear las cuentas', 2)
paso('Abra Tesorería y cree Banco de prueba con saldo inicial de $1.000.000.')
paso('Compruebe la existencia de Efectivo general.')
paso('Registre un aporte de $300.000 en Efectivo general.')
paso('Transfiera $100.000 desde Banco de prueba hacia Efectivo general.')
tabla(['Cuenta', 'Saldo esperado'], [('Banco de prueba', '$900.000'), ('Efectivo general', '$400.000'), ('Disponible total', '$1.300.000')], [3.4, 3.1])
for texto in ['La transferencia generó un egreso y un ingreso relacionados.', 'El total disponible no cambió por mover dinero entre cuentas.', 'Los movimientos muestran fecha en formato dd-mm-aaaa.']: check(texto)

heading('5. Maestros de prueba')
heading('5.1 Cliente con crédito', 2)
paso('Cree Cliente Prueba Cuenta Corriente con documento 99999992.')
paso('Habilite crédito por $200.000 y configure vencimiento a 30 días.')
heading('5.2 Proveedor', 2)
paso('Cree Distribuidora Prueba y complete los datos de identificación disponibles.')
heading('5.3 Empleado', 2)
paso('Cree el legajo PRUEBA-001, documento 99999993, cargo Auxiliar.')
paso('Seleccione modalidad quincenal y sueldo base de $300.000.')
for texto in ['Los tres registros aparecen activos.', 'Los modales se cierran al guardar.', 'Editar abre el formulario por delante del detalle.']: check(texto)

pagina(); heading('6. Inventario y carga inicial')
heading('6.1 Buscar y ajustar un producto', 2)
paso('Abra Inventario y busque por nombre o código de barras.')
paso('Ingrese la cantidad física contada usando números enteros para productos no pesables.')
paso('Use como motivo Carga inicial de existencias, o reemplace el texto al enfocar el campo.')
paso('Registre el ajuste.')
for texto in ['El stock se actualiza.', 'El foco vuelve al buscador.', 'El contenido anterior queda seleccionado para poder escanear otro código.', 'Un código inexistente también queda seleccionado para ser reemplazado.']: check(texto)
heading('6.2 Controles de inventario', 2)
for texto in ['Filtrar por categoría y marca.', 'Activar Sólo bajo mínimo.', 'Confirmar que algunos productos queden deliberadamente por debajo del mínimo.', 'Verificar Disponible, Reservado y Mínimo.']: check(texto)
nota('Conceptos.', 'Disponible es la existencia utilizable. Reservado representa unidades comprometidas. Mínimo es el nivel configurado para advertir reposición.')

pagina(); heading('7. Orden de compra y recepción')
heading('7.1 Crear la orden #1', 2)
paso('Abra Compras y seleccione Nueva orden.')
paso('Elija Distribuidora Prueba.')
paso('Busque productos por nombre o código; use Flecha abajo y Enter para agregarlos.')
paso('Agregue 6 productos, complete cantidades y costos, y guarde el borrador.')
paso('Compruebe un total de $303.956,98 y envíe la orden.')
heading('7.2 Recepción parcial', 2)
paso('Seleccione Recibir sobre la orden enviada.')
paso('Para algunos productos ingrese menos cantidad que la pedida; para otros reciba el total y deje al menos uno pendiente.')
paso('Confirme mediante el modal.')
for texto in ['La orden queda Parcial.', 'El stock aumenta solamente por lo recibido.', 'Las cantidades pendientes continúan visibles.', 'Un producto puede dejar de estar bajo mínimo si recibió suficiente cantidad.']: check(texto)
heading('7.3 Completar la recepción', 2)
paso('Vuelva a abrir la recepción y registre las cantidades restantes.')
paso('Compruebe que la orden quede Recibida.')

pagina(); heading('8. Costos y precios durante la recepción')
doc.add_paragraph('Caso validado con el producto de código 7792900000138.')
tabla(['Dato', 'Antes', 'Costo recibido', 'Después'], [('Precio de costo', '$1.399', '$1.500', '$1.500'), ('Margen', '30 %', '30 %', '30 %'), ('Precio de venta', '$1.820', 'Cálculo automático', '$1.950')], [2.0, 1.4, 1.55, 1.55])
paso('En la orden, cambie el costo unitario de $1.399 a $1.500.')
paso('Reciba el producto.')
paso('Abra Catálogo y verifique costo $1.500 y venta $1.950.')
nota('Precio manual.', 'Si un producto tiene activada la opción Precio manual, la recepción actualiza el costo pero conserva el precio de venta definido por el usuario.')

heading('9. Factura del proveedor')
paso('Abra Proveedores, seleccione Distribuidora Prueba y entre a Cuenta.')
paso('Seleccione Nueva factura.')
tabla(['Campo', 'Valor de prueba'], [('Orden relacionada', '#1'), ('Tipo', 'Factura'), ('Número', 'A-0001-00000001'), ('Emisión', 'Fecha de la prueba'), ('Vencimiento', '30 días después'), ('Total', '$303.956,98'), ('Observaciones', 'Factura de prueba vinculada a la orden #1')], [2.1, 4.4])
for texto in ['La factura queda Pendiente.', 'Saldo proveedor: $303.956,98.', 'Inicio y Tesorería muestran la deuda.', 'La orden #1 deja de aparecer en nuevas facturas.', 'El backend impide relacionar la misma orden por segunda vez.']: check(texto)

pagina(); heading('10. Pago parcial al proveedor')
paso('Abra la cuenta de Distribuidora Prueba y seleccione Registrar pago.')
tabla(['Campo', 'Valor'], [('Importe', '$100.000'), ('Medio', 'Transferencia'), ('Cuenta de Tesorería', 'Banco de prueba'), ('Referencia', 'TRANSFERENCIA-PRUEBA-001')], [2.5, 4.0])
paso('Confirme el pago y abra el comprobante.')
tabla(['Control', 'Resultado esperado'], [('Factura', 'Estado Parcial · saldo $203.956,98'), ('Banco de prueba', '$800.000'), ('Efectivo general', '$400.000'), ('Disponible total', '$1.200.000'), ('Disponible neto proyectado', '$996.043,02'), ('Libro de Tesorería', 'Egreso $100.000 · categoría proveedores')], [3.0, 3.5])
nota('Efectivo en compras directas.', 'Si la mercadería se paga al retirarla del mayorista, el proveedor sigue siendo el comercio vendedor. El pago puede salir de Efectivo general o de una caja abierta; debe seleccionarse el origen real.')

pagina(); heading('11. Gastos y servicios recurrentes')
heading('11.1 Crear el gasto', 2)
paso('Abra Gastos y servicios y seleccione Nuevo gasto.')
tabla(['Campo', 'Valor de prueba'], [('Concepto', 'Servicio de energía eléctrica - prueba'), ('Categoría', 'Energía eléctrica o equivalente'), ('Proveedor', 'Sin asociar'), ('Comprobante', 'ENERGIA-0001'), ('Total', '$120.000'), ('Vencimiento', '7 días después'), ('Recurrente', 'Sí · frecuencia mensual')], [2.25, 4.25])
for texto in ['Estado Pendiente.', 'Inicio y Tesorería muestran gastos pendientes por $120.000.', 'El dinero disponible no cambia hasta registrar el pago.']: check(texto)
heading('11.2 Preparar el próximo período', 2)
paso('Abra el gasto y seleccione Generar próximo período.')
paso('En el modal editable cambie el total sugerido a $135.000 y ajuste el vencimiento.')
paso('Guarde y confirme que existen dos períodos independientes.')
nota('Variaciones reales.', 'La frecuencia sólo propone el próximo período. El importe, la emisión, el vencimiento y el comprobante pueden cambiar porque los servicios no siempre mantienen las mismas condiciones.')
heading('11.3 Editar y anular', 2)
paso('Abra el segundo período y seleccione Editar para corregir cualquier dato mientras no tenga pagos.')
paso('Seleccione Anular e ingrese: Período generado solamente para prueba.')
for texto in ['El período anulado aparece en el filtro Anulados.', 'Su saldo pendiente queda en cero.', 'No suma en Inicio, Tesorería ni reportes.', 'Se conserva el motivo para auditoría.']: check(texto)

pagina(); heading('12. Pago parcial de un gasto')
paso('Abra el gasto original de $120.000 y seleccione Registrar pago.')
tabla(['Campo', 'Valor'], [('Importe', '$60.000'), ('Medio', 'Transferencia'), ('Cuenta de Tesorería', 'Banco de prueba'), ('Referencia', 'PAGO-ENERGIA-PRUEBA-001')], [2.5, 4.0])
paso('Confirme el pago.')
tabla(['Control', 'Resultado esperado'], [('Gasto', 'Estado Parcial · saldo $60.000'), ('Banco de prueba', '$740.000'), ('Efectivo general', '$400.000'), ('Disponible total', '$1.140.000'), ('Gastos pendientes', '$60.000'), ('Deuda proveedores', '$203.956,98'), ('Disponible neto proyectado', '$876.043,02'), ('Libro de Tesorería', 'Egreso $60.000 · categoría gastos')], [3.0, 3.5])
nota('Control corregido durante la prueba.', 'Un pago exactamente igual a la mitad dejó inicialmente el gasto oculto por un estado incorrecto. La versión actual conserva saldo $60.000 y estado Parcial; la migración 030 reparó el registro sin alterar el dinero.')

heading('12.1 Cancelar el saldo desde Efectivo general', 2)
paso('Abra nuevamente el gasto original y seleccione Registrar pago.')
tabla(['Campo', 'Valor'], [('Importe', '$60.000'), ('Medio', 'Efectivo'), ('Origen del efectivo', 'Cuenta de efectivo de Tesorería'), ('Cuenta', 'Efectivo general'), ('Referencia', 'EFECTIVO-ENERGIA-PRUEBA-002')], [2.5, 4.0])
paso('Confirme el pago y compruebe que no afecte ninguna caja abierta.')
tabla(['Control', 'Resultado validado'], [('Gasto', 'Pagado · saldo $0'), ('Banco de prueba', '$740.000'), ('Efectivo general', '$340.000'), ('Disponible total', '$1.080.000'), ('Gastos pendientes', '$0'), ('Deuda proveedores', '$203.956,98'), ('Disponible neto proyectado', '$876.043,02'), ('Libro de Tesorería', 'Egreso $60.000 · categoría gastos · Efectivo general')], [3.0, 3.5])
nota('Resultado.', 'Circuito aprobado. El gasto pasó a Pagado, el efectivo se descontó de la cuenta seleccionada y la caja operativa no fue modificada.')

heading('13. Compra directa en un mayorista')
doc.add_paragraph('Cuando una persona del supermercado retira y paga mercadería en un mayorista, el circuito conserva los mismos documentos para no perder stock, costo, comprobante ni origen del dinero.')
paso('Cree una orden para el mayorista que realmente vende la mercadería.')
paso('Cargue productos, cantidades y costos del ticket o factura.')
paso('Reciba la mercadería cuando ingresa al local.')
paso('Registre el comprobante vinculado a la orden.')
paso('Registre el pago total y seleccione el origen real: banco, Efectivo general o caja abierta.')
nota('No duplicar.', 'No registre además un egreso manual en Tesorería. Los pagos vinculados generan automáticamente el movimiento correspondiente.')

pagina(); heading('14. Controles transversales validados')
for texto in ['Altas y ediciones mediante modales; sin alertas del navegador.', 'Botones con texto de peso normal y contraste visible al pasar el mouse.', 'Búsquedas dinámicas sin botones Buscar o Limpiar innecesarios.', 'Navegación con teclado para seleccionar productos.', 'Campos numéricos seleccionan el contenido al recibir foco.', 'Fechas dd-mm-aaaa y horas de 24 horas.', 'Pagos transaccionales: si la cuenta no existe o no tiene saldo, no cambia la deuda.', 'Trazabilidad mediante comprobantes, anulaciones y movimientos automáticos.']: check(texto)

heading('15. Registro incremental de pruebas')
tabla(['Fecha', 'Circuito', 'Resultado', 'Observación'], [('10-08-2026', 'Recepción parcial y precios', 'Aprobado', 'Costo y venta actualizados.'), ('10-08-2026', 'Factura y pago proveedor', 'Aprobado', 'Tesorería integrada.'), ('10-08-2026', 'Gasto recurrente', 'Aprobado', 'Edición y anulación disponibles.'), ('10-08-2026', 'Pago parcial de gasto', 'Aprobado', 'Estado parcial corregido.'), ('10-08-2026', 'Pago desde Efectivo general', 'Aprobado', 'Gasto cancelado sin afectar caja.')], [1.1, 2.1, 1.25, 2.05])
heading('15.1 Próximos circuitos a incorporar', 2)
for texto in ['Venta contado y cierre de caja.', 'Venta a cuenta corriente y cobranza.', 'Cambio o devolución con diferencia.', 'Sueldos, adelantos y pagos desde Tesorería.', 'Cierre diario y reportes.']: check(texto)
nota('Mantenimiento del documento.', 'Después de cada circuito aprobado se actualizarán la versión, los resultados esperados y el registro de pruebas. Las capturas definitivas se incorporarán cuando la interfaz quede cerrada.')

doc.core_properties.title = 'Guía de Prueba Funcional - Sistema de Gestión La 91 Supermercado'
doc.core_properties.subject = 'Recorrido incremental de capacitación y aceptación'
doc.core_properties.author = 'La 91 Supermercado'
doc.core_properties.keywords = 'supermercado, prueba funcional, instructivo, capacitación, aceptación'
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
